import io
from docx import Document


async def extract_text_from_docx(stream: io.BytesIO) -> str:
    """
    Extract all text from a .docx invoice file.
    Args: stream (io.BytesIO): Stream of the .docx file.
    Returns: str: Extracted and cleaned text from the document.
    """
    try:
        # Load the document from memory stream
        doc = Document(stream)
        full_text = []

        # Extract text from paragraphs
        full_text.extend(
            para.text.strip()
            for para in doc.paragraphs
            if para.text.strip()
        )

        # Extract text from all tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)

    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")
